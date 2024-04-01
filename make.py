import subprocess

from docx import Document as docxDocument
from docx.oxml.ns import qn
from docx2pdf import convert


font_name = "宋体"
tmp_docx = "tmp.docx"

# `article` dict should at least contains `title`, `date` and `content`, and `article["content"]` should be list[str] type.
def make_docx(docx_path: str, articles: list[dict]):
    doc = docxDocument()
    for arti in articles:
        doc.add_heading(arti["title"], level=1)
        doc.add_paragraph(arti["date"])
        for para in arti["content"]:
            doc_para = doc.add_paragraph(para)
        # modify the font of all `run` objects of the current `arti`.
        for doc_para in doc.paragraphs:
            for doc_run in doc_para.runs:
                doc_run.font.name = font_name
                doc_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    doc.save(docx_path)

def make_pdf(output_path: str, articles: list):
    # make a temporary .docx file first, then convert it to .pdf.
    make_docx(tmp_docx, articles)
    convert(tmp_docx, output_path)
    subprocess.run(["rm", tmp_docx])
