import os
import jieba

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX


sensitive_wordlist_path = "./sensitive_wordlist.txt"
output_dir="outputs/"
font_name = "宋体"

class Marker:
    def __init__(self) -> None:
        self.wordlist = set()
        with open(sensitive_wordlist_path, "r", encoding="UTF-8") as txtfile:
            self.wordlist = set([line.strip() for line in txtfile.readlines()])

    def search(self, word: str) -> bool:
        return word in self.wordlist


def mark(wordmarking_path: str):
    doc = Document(wordmarking_path)
    doc_marked = Document()
    marker = Marker()
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            para_marked = doc_marked.add_heading("")
            para_marked.style = para.style
        else:
            para_marked = doc_marked.add_paragraph("")
        for word in jieba.cut(para.text):
            run = para_marked.add_run(word)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            if marker.search(word):
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        para_marked.add_run("\n")
    
    doc_marked.save(output_dir + os.path.split(wordmarking_path)[-1].split(".")[0] + "_marked.docx")

